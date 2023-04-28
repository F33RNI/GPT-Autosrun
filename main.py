"""
 Copyright (C) 2022 Fern Lane, GPT-Autosrun project
 Licensed under the GNU Affero General Public License, Version 3.0 (the "License");
 you may not use this file except in compliance with the License.
 You may obtain a copy of the License at
       https://www.gnu.org/licenses/agpl-3.0.en.html
 Unless required by applicable law or agreed to in writing, software
 distributed under the License is distributed on an "AS IS" BASIS,
 WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
 See the License for the specific language governing permissions and
 limitations under the License.
 IN NO EVENT SHALL THE AUTHOR BE LIABLE FOR ANY CLAIM, DAMAGES OR
 OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE,
 ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR
 OTHER DEALINGS IN THE SOFTWARE.
"""

import os
import shutil
import time
import uuid

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Mm

from revChatGPT.V3 import Chatbot

# Group and name
GROUP_REPLACE_WITH = '224-371'
NAME_REPLACE_WITH = 'Фмилиия И.О.'

# Output file ({0} - task number)
OUTPUT_FILE_FORMAT = '224-371_Name_Surname_NITvNiPD_Pr{0}.docx'

# Go to https://platform.openai.com/account/api-keys and generate new API key and paste it below
OPENAI_API_KEY = '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'

# Formatting settings
PARAGRAPH_TASK = 'Задание на практическую работу'
PARAGRAPH_ANSWERS = 'Ход работы'
PARAGRAPH_SOURCES = 'Список литературы'
PARAGRAPH_FONT_SIZE_PT = 14
PARAGRAPH_FONT_NAME = 'Times New Roman'
PARAGRAPH_LEFT_INDENT_MM = 12.5
PARAGRAPH_LINE_SPACING = WD_LINE_SPACING.ONE_POINT_FIVE

# Ask ChatGPT to write sources?
GENERATE_SOURCES = False

# Skip files with this name
SKIP_TASKS = []

# Replaces parts of the text in request to chatGPT
GPT_REQUEST_REPLACE_FROM = ['Вашей специальностью']
GPT_REQUEST_REPLACE_TO = ['разработкой мобильных приложений']

# Requests for chatGPT
LECTURE_REQUEST_LIMIT = 3999
LECTURE_REQUEST_FIRST = 'Я напишу тебе лекцию в несколько сообщений.' \
                        ' Потом попрошу ответить на некоторые вопросы, ' \
                        'используя текст этой лекции. Пока что я буду тебе писать части одной лекции.' \
                        ' Когда нужно будет ответить на вопросы я напишу об этом. ' \
                        'Поэтому пока просто "принял части лекций". Без уточнений. Вот первая часть лекции:\n'

LECTURE_REQUEST_OTHER = 'Продолжение лекции:\n'

QUESTION_REQUEST = 'Теперь, частично используя лекцию напиши максимально подробно на тему:\n{0}\n' \
                   'Ответь без личных вводных слов. Сразу напиши максимальный подробный ответ на вопрос'

REQUEST_SOURCES = 'Напиши по-русски и оформи список литературы из книг с номерами ' \
                  'страниц или интернет источников по теме:\n{0}'

REQUEST_CONTINUE = 'Продолжай'

# How many 'continue' requests can be used for each question
REQUEST_CONTINUE_MAX_TIMES = 2

# What to replace in title page
TASK_N = 'TASKN'
TASK_TOPIC = 'TASKTOPIC'
TASK_GROUP = 'GROUP'
TASK_NAME = 'NAME'

# Files
LECTURES_DIR = 'lectures'
TOPICS_FILE = 'topics.txt'
SUB_TASKS_DIR = 'sub_tasks'
MAIN_TASKS_DIR = 'main_tasks'
RESULT_DIR = 'output'
TITLE_PAGE_FILE = 'title.docx'

# Rate limiting
MIN_TIME_BETWEEN_REQUESTS_SECONDS = 30

# Too many requests in 1 hour or Rate limit reached handling
TOO_MANY_REQUESTS_EXCEPTION = 'Too many requests in 1 hour'
RATE_LIMIT_REACHED_EXCEPTION = 'Rate limit reached'
TOO_MANY_REQUESTS_WAIT_SECONDS = 600
RATE_LIMIT_REACHED_WAIT_SECONDS = 30


def replace_text_in_paragraph(paragraph_, key, value):
    """
    Replaces text in paragraph without touching style
    :param paragraph_:
    :param key:
    :param value:
    :return:
    """
    if key in paragraph_.text:
        inline = paragraph_.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def format_lines(lines, remove_empty_lines=False, remove_ending=False):
    """
    Processes lines from file
    :param lines:
    :param remove_empty_lines:
    :param remove_ending:
    :return:
    """
    if remove_empty_lines:
        lines = [line for line in lines if len(line.replace('\n', '').strip()) > 1]

    for line_n in range(len(lines)):
        lines[line_n] = lines[line_n].strip()
        while '  ' in lines[line_n]:
            lines[line_n] = lines[line_n].replace('  ', ' ')

        if remove_ending:
            if lines[line_n].endswith('.'):
                lines[line_n] = lines[line_n][: -1]

            if lines[line_n].endswith(';'):
                lines[line_n] = lines[line_n][: -1]

    return lines


def document_add_header(document_, header_name: str, page_break=False):
    """
    Adds center header to the document
    :param document_:
    :param header_name:
    :param page_break:
    :return:
    """
    if page_break:
        document_.add_page_break()
    paragraph_ = document_.add_paragraph()
    paragraph_.alignment = 1
    paragraph_run_ = paragraph_.add_run(header_name)
    paragraph_run_.font.size = Pt(PARAGRAPH_FONT_SIZE_PT)
    paragraph_run_.font.name = PARAGRAPH_FONT_NAME
    paragraph_run_.bold = True


def document_add_paragraph(document_, item_text: str, is_list=False, indent=False, bold=False, justify=False):
    """
    Adds paragraph (or list) to the document
    :param document_:
    :param item_text:
    :param is_list:
    :param indent:
    :param bold:
    :param justify:
    :return:
    """
    paragraph_ = document_.add_paragraph()
    paragraph_.alignment = 3 if justify else 0
    paragraph_format_ = paragraph_.paragraph_format
    paragraph_format_.line_spacing_rule = PARAGRAPH_LINE_SPACING
    if indent:
        paragraph_format_.first_line_indent = Mm(PARAGRAPH_LEFT_INDENT_MM)

    if is_list:
        item_text_splitted = item_text.split('. ')
        paragraph_run_ = paragraph_.add_run(
            item_text_splitted[0].strip() + '.\t' + '. '.join(item_text_splitted[1:]).strip())
    else:
        paragraph_run_ = paragraph_.add_run(item_text)

    paragraph_run_.font.size = Pt(PARAGRAPH_FONT_SIZE_PT)
    paragraph_run_.font.name = PARAGRAPH_FONT_NAME
    if bold:
        paragraph_run_.bold = True


def ask_chatbot(chatbot, request, conversation_id, last_request_time):
    """
    Asks chatGPT
    :param chatbot:
    :param request:
    :param conversation_id:
    :param last_request_time:
    :return:
    """
    # Wait
    if (time.time() - last_request_time) < MIN_TIME_BETWEEN_REQUESTS_SECONDS:
        print('Waiting before next request...')
        while (time.time() - last_request_time) < MIN_TIME_BETWEEN_REQUESTS_SECONDS:
            time.sleep(1)

    response = ''
    print('Asking: ' + request)

    # Save request time
    last_request_time = time.time()

    try:
        for data in chatbot.ask_stream(request, convo_id=conversation_id):
            # Get response
            response += data
            print(data, flush=True, end="")

    except Exception as e:
        print('Error! ' + str(e))

        # Too many requests in 1 hour
        if TOO_MANY_REQUESTS_EXCEPTION in str(e):
            print('Waiting ' + str(TOO_MANY_REQUESTS_WAIT_SECONDS) + ' seconds...')
            time.sleep(TOO_MANY_REQUESTS_WAIT_SECONDS)

            # Ask again
            return ask_chatbot(chatbot, request, conversation_id, last_request_time)

        elif RATE_LIMIT_REACHED_EXCEPTION in str(e):
            print('Waiting ' + str(RATE_LIMIT_REACHED_WAIT_SECONDS) + ' seconds...')
            time.sleep(RATE_LIMIT_REACHED_WAIT_SECONDS)

            # Ask again
            return ask_chatbot(chatbot, request, conversation_id, last_request_time)

    # Check response length
    if len(response.split(' ')) > 1:
        # Return response
        print('Responce words: ' + str(len(response.split(' '))))
        return response, conversation_id, last_request_time

    # No words
    else:
        print('Empty response! Asking again')
        return ask_chatbot(chatbot, request, conversation_id, last_request_time)


def main():
    # Initialize chatGPT
    chatbot = Chatbot(api_key=OPENAI_API_KEY)

    # Create output dir
    if not os.path.exists(RESULT_DIR):
        shutil.rmtree(RESULT_DIR, ignore_errors=True)
        os.makedirs(RESULT_DIR)

    # Read topics
    topics_file = open(TOPICS_FILE, 'r', encoding='utf-8')
    topics = format_lines(topics_file.readlines(), remove_empty_lines=False, remove_ending=True)
    topics_file.close()
    print('Topics: ' + str(topics))

    # Variable to store time between requests
    last_request_time = 0

    # List all files in Tasks directory
    for file in os.listdir(SUB_TASKS_DIR):
        # Check if it is txt
        if file.endswith('.txt') and len(file.split('.')) == 2 and int(file.split('.')[0]) > 0:
            # Get task number and index
            task_number = int(file.split('.')[0])
            task_index = task_number - 1

            if task_number not in SKIP_TASKS:
                print('Processing task n ' + str(task_number))

                # Read files
                main_task_file = open(os.path.join(MAIN_TASKS_DIR, file), 'r', encoding='utf-8')
                main_task_lines = format_lines(main_task_file.readlines(), remove_empty_lines=True, remove_ending=True)
                main_task_file.close()

                sub_task_file = open(os.path.join(SUB_TASKS_DIR, file), 'r', encoding='utf-8')
                sub_task_lines = format_lines(sub_task_file.readlines(), remove_empty_lines=True, remove_ending=True)
                sub_task_file.close()

                # Read lecture
                lecture_file = open(os.path.join(LECTURES_DIR, file), 'r', encoding='utf-8')
                lecture_lines = format_lines(lecture_file.readlines(), remove_empty_lines=True, remove_ending=False)
                lecture_file.close()

                # Copy title page
                output_filename = os.path.join(RESULT_DIR, OUTPUT_FILE_FORMAT.format(task_number))
                shutil.copyfile(TITLE_PAGE_FILE, output_filename)

                # Start new document
                document = Document(output_filename)

                # Replace title page fields
                for paragraph in document.paragraphs:
                    replace_text_in_paragraph(paragraph, TASK_N, str(task_number))
                    replace_text_in_paragraph(paragraph, TASK_TOPIC, topics[task_index])
                    replace_text_in_paragraph(paragraph, TASK_NAME, NAME_REPLACE_WITH)
                    replace_text_in_paragraph(paragraph, TASK_GROUP, GROUP_REPLACE_WITH)

                # Add main tasks
                document_add_header(document, PARAGRAPH_TASK, page_break=True)
                for i in range(len(main_task_lines)):
                    document_add_paragraph(document,
                                           main_task_lines[i] + (';' if i < (len(main_task_lines) - 1) else '.'),
                                           is_list=True, justify=True)
                document.add_paragraph()

                # Start each task with new conversation_id
                conversation_id = str(uuid.uuid4())

                # Send lecture
                lecture_request = ""

                while True:
                    if len(lecture_lines) <= 0:
                        break

                    lecture_request = LECTURE_REQUEST_FIRST if lecture_request == "" else LECTURE_REQUEST_OTHER
                    while len(lecture_lines) > 0 and len(lecture_request) < LECTURE_REQUEST_LIMIT:
                        last_pop_ = lecture_lines.pop(0)
                        if len(lecture_request + last_pop_) <= LECTURE_REQUEST_LIMIT:
                            lecture_request += last_pop_
                        else:
                            lecture_lines.insert(0, last_pop_)
                            break

                    response, conversation_id, last_request_time = ask_chatbot(chatbot,
                                                                               lecture_request,
                                                                               conversation_id,
                                                                               last_request_time)
                    print("", flush=True)

                # Add sub-tasks
                document_add_header(document, PARAGRAPH_ANSWERS, page_break=False)
                for i in range(len(sub_task_lines)):
                    print("", flush=True)
                    print('Processing sub-task ' + str(i + 1) + '/' + str(len(sub_task_lines))
                          + ': ' + sub_task_lines[i])

                    # Add paragraph to doc
                    document_add_paragraph(document, sub_task_lines[i], is_list=True, bold=True, justify=True)

                    # Replace text before requesting
                    request = '.'.join(sub_task_lines[i].split('.')[1:])
                    for replace_gpt_i in range(len(GPT_REQUEST_REPLACE_FROM)):
                        request = request.replace(GPT_REQUEST_REPLACE_FROM[replace_gpt_i],
                                                  GPT_REQUEST_REPLACE_TO[replace_gpt_i])

                    # Format request
                    request = QUESTION_REQUEST.format(request)

                    # Ask chatGPT
                    response, conversation_id, last_request_time = ask_chatbot(chatbot,
                                                                               request,
                                                                               conversation_id,
                                                                               last_request_time)

                    # Split into lines
                    response_lines = format_lines(response.split('\n'), remove_empty_lines=True, remove_ending=False)

                    response_check_words_last = ''
                    response_lines_temp = []
                    response_continue_counter = 0
                    while True:
                        # Print response ending
                        response_combined = ' '.join(response_lines + response_lines_temp)
                        response_words = response_combined.split(' ')
                        response_words = [word for word in response_words if len(word.strip()) > 0]

                        if len(response_words) > 2:
                            response_words = response_words[-2:]
                        response_check_words = ' '.join(response_words)
                        print('Ends with: ...' + response_check_words)
                        # Ends with dot -> done
                        if response_words[-1][-1] == '.':
                            print('Ends with dot. Finishing question...')
                            response_lines += response_lines_temp
                            break

                        # Same response -> done
                        elif response_check_words == response_check_words_last:
                            print('Same words. Finishing question...')
                            break

                        # Try to continue
                        else:
                            print('Continuing... Continue task n ' + str(response_continue_counter + 1))
                            response_continue_counter += 1
                            response, conversation_id, last_request_time = ask_chatbot(chatbot,
                                                                                       REQUEST_CONTINUE,
                                                                                       conversation_id,
                                                                                       last_request_time)
                            response_lines += response_lines_temp
                            response_lines_temp = []
                            response_lines_temp += format_lines(response.split('\n'),
                                                                remove_empty_lines=True, remove_ending=True)

                            # Stop continue requests
                            if response_continue_counter >= REQUEST_CONTINUE_MAX_TIMES:
                                print('Max continue requests achieved. Finishing question...')
                                break

                        # Save for next check cycle
                        response_check_words_last = response_check_words

                    # Add response
                    for response_line in response_lines:
                        document_add_paragraph(document, response_line, indent=True, justify=True)
                    document.add_paragraph()

                # Add sources
                if GENERATE_SOURCES:
                    document_add_header(document, PARAGRAPH_SOURCES, page_break=True)
                    request = REQUEST_SOURCES.format(topics[task_index])
                    response, conversation_id, last_request_time = ask_chatbot(chatbot,
                                                                               request,
                                                                               'default',
                                                                               last_request_time)

                    # Split into lines
                    response_lines = format_lines(response.split('\n'), remove_empty_lines=True, remove_ending=True)

                    # Add to the document
                    for i in range(len(response_lines)):
                        document_add_paragraph(document,
                                               response_lines[i] + (';' if i < (len(main_task_lines) - 1) else '.'),
                                               indent=True, is_list=True, justify=True)
                else:
                    print('Skipping sources')

                # Save document
                document.save(output_filename)
                print('Document: ' + str(output_filename) + ' saved!')

            # Skip task
            else:
                print('Skipping task n ' + str(task_number))


if __name__ == '__main__':
    main()
